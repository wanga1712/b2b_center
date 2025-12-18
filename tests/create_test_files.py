"""
Скрипт для создания тестовых файлов различных форматов.

Создает тестовые файлы:
- Excel (.xlsx, .xls)
- Word (.docx)
- PDF (обычный текст)
- PDF (отсканированный - симуляция через изображение)

Запуск:
    python tests/create_test_files.py
"""

from pathlib import Path
import os

# Создаем директорию для тестовых файлов
TEST_FILES_DIR = Path(__file__).parent / "test_files"
TEST_FILES_DIR.mkdir(exist_ok=True)


def create_test_excel_xlsx():
    """Создает тестовый Excel файл .xlsx"""
    try:
        import openpyxl
        
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Тестовый лист"
        
        # Заголовки
        ws['A1'] = "Товар"
        ws['B1'] = "Количество"
        ws['C1'] = "Цена"
        ws['D1'] = "Сумма"
        
        # Данные
        test_data = [
            ["Контейнер мусорный 240л", 10, 1500, 15000],
            ["Мешок для мусора 120л", 50, 25, 1250],
            ["Пакет полиэтиленовый", 100, 5, 500],
            ["Крышка для контейнера", 5, 300, 1500],
        ]
        
        for row_idx, row_data in enumerate(test_data, start=2):
            for col_idx, value in enumerate(row_data, start=1):
                ws.cell(row=row_idx, column=col_idx, value=value)
        
        file_path = TEST_FILES_DIR / "test_excel.xlsx"
        wb.save(file_path)
        print(f"✅ Создан тестовый Excel файл: {file_path}")
        return file_path
        
    except ImportError:
        print("⚠️  openpyxl не установлен, пропускаем создание .xlsx файла")
        return None


def create_test_excel_xls():
    """Создает тестовый Excel файл .xls (старый формат)"""
    try:
        import xlwt
        
        wb = xlwt.Workbook()
        ws = wb.add_sheet("Тестовый лист")
        
        # Заголовки
        ws.write(0, 0, "Товар")
        ws.write(0, 1, "Количество")
        ws.write(0, 2, "Цена")
        ws.write(0, 3, "Сумма")
        
        # Данные
        test_data = [
            ["Контейнер мусорный 240л", 10, 1500, 15000],
            ["Мешок для мусора 120л", 50, 25, 1250],
            ["Пакет полиэтиленовый", 100, 5, 500],
        ]
        
        for row_idx, row_data in enumerate(test_data, start=1):
            for col_idx, value in enumerate(row_data):
                ws.write(row_idx, col_idx, value)
        
        file_path = TEST_FILES_DIR / "test_excel.xls"
        wb.save(file_path)
        print(f"✅ Создан тестовый Excel файл .xls: {file_path}")
        return file_path
        
    except ImportError:
        print("⚠️  xlwt не установлен, пропускаем создание .xls файла")
        return None


def create_test_word():
    """Создает тестовый Word документ .docx"""
    try:
        import docx
        
        doc = docx.Document()
        
        # Заголовок
        doc.add_heading('Тестовый документ Word', 0)
        
        # Параграфы
        doc.add_paragraph('Это тестовый документ для проверки парсинга Word файлов.')
        doc.add_paragraph('')
        doc.add_paragraph('Список товаров:')
        
        # Список
        items = [
            'Контейнер мусорный 240л - 10 шт',
            'Мешок для мусора 120л - 50 шт',
            'Пакет полиэтиленовый - 100 шт',
        ]
        for item in items:
            doc.add_paragraph(item, style='List Bullet')
        
        # Таблица
        doc.add_paragraph('')
        table = doc.add_table(rows=4, cols=3)
        table.style = 'Light Grid Accent 1'
        
        # Заголовки таблицы
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Товар'
        header_cells[1].text = 'Количество'
        header_cells[2].text = 'Цена'
        
        # Данные таблицы
        table_data = [
            ['Контейнер мусорный 240л', '10', '1500'],
            ['Мешок для мусора 120л', '50', '25'],
            ['Пакет полиэтиленовый', '100', '5'],
        ]
        
        for row_idx, row_data in enumerate(table_data, start=1):
            row_cells = table.rows[row_idx].cells
            for col_idx, value in enumerate(row_data):
                row_cells[col_idx].text = value
        
        file_path = TEST_FILES_DIR / "test_word.docx"
        doc.save(file_path)
        print(f"✅ Создан тестовый Word документ: {file_path}")
        return file_path
        
    except ImportError:
        print("⚠️  python-docx не установлен, пропускаем создание Word файла")
        return None


def create_test_pdf():
    """Создает тестовый PDF файл с текстом"""
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        
        file_path = TEST_FILES_DIR / "test_pdf.pdf"
        c = canvas.Canvas(str(file_path), pagesize=letter)
        width, height = letter
        
        # Заголовок
        c.setFont("Helvetica-Bold", 16)
        c.drawString(100, height - 50, "Тестовый PDF документ")
        
        # Текст
        c.setFont("Helvetica", 12)
        y_position = height - 100
        
        lines = [
            "Это тестовый PDF документ для проверки парсинга.",
            "",
            "Список товаров:",
            "1. Контейнер мусорный 240л - 10 шт - 1500 руб",
            "2. Мешок для мусора 120л - 50 шт - 25 руб",
            "3. Пакет полиэтиленовый - 100 шт - 5 руб",
            "",
            "Общая сумма: 16850 руб",
        ]
        
        for line in lines:
            c.drawString(100, y_position, line)
            y_position -= 20
        
        c.save()
        print(f"✅ Создан тестовый PDF файл: {file_path}")
        return file_path
        
    except ImportError:
        print("⚠️  reportlab не установлен, пропускаем создание PDF файла")
        return None


def create_test_pdf_image():
    """Создает PDF с изображением (имитация отсканированного)"""
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        from PIL import Image, ImageDraw, ImageFont
        import io
        
        # Создаем изображение с текстом
        img_width, img_height = 800, 600
        img = Image.new('RGB', (img_width, img_height), color='white')
        draw = ImageDraw.Draw(img)
        
        # Пробуем использовать шрифт
        try:
            # Windows шрифт
            font = ImageFont.truetype("arial.ttf", 24)
        except:
            try:
                font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 24)
            except:
                font = ImageFont.load_default()
        
        # Рисуем текст на изображении
        text_lines = [
            "Тестовый отсканированный PDF",
            "",
            "Товар: Контейнер мусорный 240л",
            "Количество: 10 шт",
            "Цена: 1500 руб",
        ]
        
        y = 50
        for line in text_lines:
            draw.text((50, y), line, fill='black', font=font)
            y += 50
        
        # Сохраняем изображение во временный файл
        temp_img_path = TEST_FILES_DIR / "temp_test_image.png"
        img.save(temp_img_path)
        
        # Создаем PDF с изображением
        file_path = TEST_FILES_DIR / "test_pdf_scanned.pdf"
        c = canvas.Canvas(str(file_path), pagesize=letter)
        c.drawImage(str(temp_img_path), 0, 0, width=letter[0], height=letter[1])
        c.save()
        
        # Удаляем временный файл
        temp_img_path.unlink()
        
        print(f"✅ Создан тестовый PDF (отсканированный): {file_path}")
        print(f"   Примечание: для реального OCR требуется установленный Tesseract")
        return file_path
        
    except ImportError:
        print("⚠️  reportlab или PIL не установлен, пропускаем создание отсканированного PDF")
        return None


def main():
    """Создает все тестовые файлы"""
    print("=" * 60)
    print("Создание тестовых файлов для проверки парсинга документов")
    print("=" * 60)
    print()
    
    files_created = []
    
    # Excel файлы
    xlsx_file = create_test_excel_xlsx()
    if xlsx_file:
        files_created.append(xlsx_file)
    
    xls_file = create_test_excel_xls()
    if xls_file:
        files_created.append(xls_file)
    
    # Word файлы
    word_file = create_test_word()
    if word_file:
        files_created.append(word_file)
    
    # PDF файлы
    pdf_file = create_test_pdf()
    if pdf_file:
        files_created.append(pdf_file)
    
    pdf_scanned_file = create_test_pdf_image()
    if pdf_scanned_file:
        files_created.append(pdf_scanned_file)
    
    print()
    print("=" * 60)
    print(f"Создано файлов: {len(files_created)}")
    print(f"Директория: {TEST_FILES_DIR}")
    print("=" * 60)
    
    if files_created:
        print("\nСозданные файлы:")
        for file_path in files_created:
            print(f"  - {file_path.name}")


if __name__ == "__main__":
    main()

